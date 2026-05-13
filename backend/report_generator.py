from docx import Document
from docx.shared import Inches

def generate_report(data, predictions):

    document = Document()

    # TITLE
    document.add_heading('CircAI Sustainability Assessment Report', 0)

    # PROJECT DETAILS
    document.add_heading('Project Information', level=1)

    document.add_paragraph(
    f"Project Name: {data['project_name']}"
)

    document.add_paragraph(
    f"Description: {data['description']}"
)

    document.add_heading('Assessment Inputs', level=1)

    document.add_paragraph(
    f"Metal Type: {data['metal_type']}"
)

    document.add_paragraph(
    f"Energy Mix: {data['energy_mix']}"
)

    document.add_paragraph(
    f"Production Route: {data['production_route']}"
)

    document.add_paragraph(
    f"Transport Distance: {data['transport_distance_km']} km"
)

    document.add_paragraph(
    f"Product Lifespan: {data['product_lifespan_years']} years"
)

    document.add_paragraph(
    f"End-of-Life Treatment: {data['end_of_life_treatment']}"
)

    # RESULTS
    document.add_heading('ML Prediction Results', level=1)

    document.add_paragraph(
        f"Carbon Footprint: {predictions['carbon_footprint']} tCO₂e"
    )

    document.add_paragraph(
        f"Circularity Score: {predictions['circularity_score']}/100"
    )

    document.add_paragraph(
        f"Sustainability Grade: {predictions['sustainability_grade']}"
    )

    # RECOMMENDATIONS
    document.add_heading('AI Recommendations', level=1)

    recommendations = []

    if predictions['carbon_footprint'] > 20:
        recommendations.append(
            "Switch to renewable energy sources to reduce emissions."
        )

    if predictions['circularity_score'] < 50:
        recommendations.append(
            "Increase recycled material usage."
        )

    if predictions['sustainability_grade'] == 'D':
        recommendations.append(
            "Redesign manufacturing workflow for sustainability."
        )

    if len(recommendations) == 0:
        recommendations.append(
            "Overall sustainability performance is satisfactory."
        )

    for rec in recommendations:
        document.add_paragraph(rec, style='List Bullet')

    # SAVE FILE
    output_path = "generated_report.docx"

    document.save(output_path)

    return output_path